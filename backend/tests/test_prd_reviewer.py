from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.main import app
from app.models import ClassificationLevel, ReadinessStatus, ReviewConfig
from app.services.prd_reviewer import review_prd
from app.models import ReviewRequest


client = TestClient(app)


def test_review_flags_missing_metrics_and_scope() -> None:
    response = review_prd(
        ReviewRequest(
            feature_name="Smart Recommendations",
            platform="Web",
            prd_text=(
                "We will launch a new recommendation capability for users. "
                "The problem is that users cannot discover relevant templates. "
                "This should increase conversion by 20%."
            ),
        )
    )

    assert response.classification_level == ClassificationLevel.FULL
    assert response.overall_assessment in {ReadinessStatus.NOT_READY, ReadinessStatus.READY_WITH_CAVEATS}
    assert any(item.criteria == "Metrics & Data" and item.evaluation == "Needs Review" for item in response.dimensional_analysis)
    assert any(finding.issue == "Unsupported headroom assumption" for finding in response.detailed_findings)
    assert len(response.detailed_findings) <= 3


def test_sensitive_prd_routes_to_advanced_scrutiny() -> None:
    response = review_prd(
        ReviewRequest(
            feature_name="Pricing Rules",
            platform="API",
            prd_text=(
                "Problem: admins need pricing policy control. "
                "Scope: add pricing rules and payment permissions. "
                "User experience must cover errors. Metrics include conversion guardrails and fraud review."
            ),
        )
    )

    assert response.classification_level == ClassificationLevel.FULL_WITH_SCRUTINY
    assert response.token_plan.route == "advanced-scrutiny"


def test_api_returns_scorecard_shape() -> None:
    result = client.post(
        "/api/reviews",
        json={
            "feature_name": "Internal Migration",
            "platform": "Backend",
            "prd_text": (
                "This migration updates an internal workflow. "
                "Scope covers moving batch jobs to the new worker. "
                "Metrics include error rate and latency guardrails. "
                "UX notes include fallback and error states."
            ),
            "token_budget": 3000,
        },
    )

    assert result.status_code == 200
    payload = result.json()
    assert payload["feature_name"] == "Internal Migration"
    assert "dimensional_analysis" in payload
    assert "token_plan" in payload
    assert payload["token_plan"]["budget_status"]


def test_config_limits_findings_and_extra_sensitive_terms() -> None:
    response = review_prd(
        ReviewRequest(
            feature_name="KYC Review",
            platform="Web",
            prd_text=(
                "Problem: operators need a new workflow for customer review. "
                "Scope is not finalized. This should increase approvals by 15%."
            ),
            config=ReviewConfig(max_findings=1, extra_sensitive_terms=["KYC"]),
        )
    )

    assert response.classification_level == ClassificationLevel.FULL_WITH_SCRUTINY
    assert len(response.detailed_findings) == 1


def test_extract_docx_endpoint() -> None:
    document = Document()
    document.add_heading("PRD: DOCX Import", level=1)
    document.add_paragraph(
        "Problem: product teams need to import PRDs from Word documents. "
        "Scope: extract readable paragraphs and tables. "
        "User Experience: imported text should populate the editor. "
        "Metrics: track import success and parser failure guardrails."
    )
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    result = client.post(
        "/api/reviews/extract-docx",
        files={
            "file": (
                "sample-prd.docx",
                buffer.getvalue(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )

    assert result.status_code == 200
    payload = result.json()
    assert payload["filename"] == "sample-prd.docx"
    assert "import PRDs from Word documents" in payload["text"]


def test_product_context_review_workflow() -> None:
    product_result = client.post(
        "/api/products",
        json={"name": "Workflow Test Product", "description": "Uses KB context before PRD review."},
    )
    assert product_result.status_code == 200
    product = product_result.json()

    context_result = client.post(
        f"/api/products/{product['id']}/contexts",
        json={
            "title": "Launch Standards",
            "text": "All PRDs must define guardrail metrics, non-goals, and adjacent system owners.",
            "source_type": "manual",
        },
    )
    assert context_result.status_code == 200

    review_result = client.post(
        f"/api/products/{product['id']}/reviews",
        json={
            "feature_name": "KB-backed Review",
            "platform": "Web",
            "prd_text": (
                "Problem: teams need PRD review with product context. "
                "Scope: use uploaded knowledge base during evaluation. "
                "UX: show blocker first and handle error states. "
                "Metrics: success rate and guardrail latency are tracked."
            ),
            "config": {"max_findings": 2},
        },
    )
    assert review_result.status_code == 200
    payload = review_result.json()
    assert payload["feature_name"] == "KB-backed Review"
    assert len(payload["detailed_findings"]) <= 2


def test_agent_hardness_loop_mode_returns_trace() -> None:
    result = client.post(
        "/api/reviews",
        json={
            "feature_name": "KYC Pricing Policy",
            "platform": "Web",
            "prd_text": (
                "Problem: operators need a new pricing policy workflow. "
                "Scope: add permissioned controls. UX: handle error states. "
                "Metrics: track conversion and latency guardrails."
            ),
            "config": {"review_mode": "agent", "max_findings": 2, "extra_sensitive_terms": ["KYC"]},
        },
    )

    assert result.status_code == 200
    payload = result.json()
    assert payload["agent_architecture"] in {
        "hardness-loop-deterministic",
        "langchain-react-hardness-loop",
        "ollama-native-tool-agent",
    }
    assert payload["agent_trace"]
    assert payload["agent_trace"][0]["tool"] == "classify_review_hardness"
